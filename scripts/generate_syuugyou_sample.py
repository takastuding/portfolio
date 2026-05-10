"""
開業したて・個人事業主・副業者・初めて従業員を雇う方向けの
就業規則サンプル雛形（Word .docx）を生成するスクリプト。

使い方:
    python3 scripts/generate_syuugyou_sample.py

出力先:
    public/samples/syuugyou-kisoku-sample.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT = Path(__file__).resolve().parent.parent / "public" / "samples" / "syuugyou-kisoku-sample.docx"


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)


def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x57, 0x57, 0x57)
    p.paragraph_format.space_after = Pt(6)


def main():
    doc = Document()

    # 余白
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 表紙
    add_title(doc, "就業規則（サンプル雛形）")
    add_para(doc, "")
    add_para(doc, "対象：開業したて・個人事業主・副業者・初めて従業員を雇う方")
    add_note(doc, "本サンプルは、これから初めて従業員を雇う方や、開業して間もない事業者の方が、就業規則の最小構成を理解し、自社の実態に合わせて調整しやすくするための雛形です。実際に運用される際は、社会保険労務士など専門家にご相談ください。")

    add_para(doc, "")
    add_para(doc, "事業所名：＿＿＿＿＿＿＿＿＿＿＿＿＿＿")
    add_para(doc, "代表者　：＿＿＿＿＿＿＿＿＿＿＿＿＿＿")
    add_para(doc, "施行日　：　　年　　月　　日")

    doc.add_page_break()

    # 第1章 総則
    add_heading(doc, "第1章　総則", level=1)

    add_heading(doc, "第1条（目的）", level=2)
    add_para(doc, "1. この規則は、○○（事業所名。以下「事業主」という。）に勤務する従業員の労働条件、服務規律その他の就業に関する事項を定めることを目的とする。")
    add_para(doc, "2. この規則に定めのない事項については、労働基準法その他の関係法令の定めるところによる。")

    add_heading(doc, "第2条（適用範囲）", level=2)
    add_para(doc, "1. この規則は、事業主に雇用されるすべての従業員（正社員、契約社員、パートタイマー、アルバイトを含む）に適用する。ただし、雇用形態によって取扱いが異なる事項は、別に定める。")
    add_note(doc, "※開業初期で正社員のみの場合は、ただし書きを削除しても構いません。")

    add_heading(doc, "第3条（規則の遵守）", level=2)
    add_para(doc, "事業主および従業員は、この規則を遵守し、誠実に職務を遂行しなければならない。")

    # 第2章 採用・労働契約
    add_heading(doc, "第2章　採用・労働契約", level=1)

    add_heading(doc, "第4条（採用）", level=2)
    add_para(doc, "事業主は、応募者の中から選考のうえ、採用を決定する。採用にあたっては、書面により労働条件を明示する。")

    add_heading(doc, "第5条（試用期間）", level=2)
    add_para(doc, "1. 新たに採用した従業員には、原則として採用日から3か月間の試用期間を設ける。")
    add_para(doc, "2. 試用期間中または試用期間満了時に、勤務態度・業務遂行能力等に照らし社員として不適格と認めたときは、本採用を行わないことがある。")
    add_note(doc, "※試用期間は1〜6か月の範囲で設定するのが一般的です。")

    # 第3章 服務規律
    add_heading(doc, "第3章　服務規律", level=1)

    add_heading(doc, "第6条（服務の基本原則）", level=2)
    add_para(doc, "従業員は、職務上の責任を自覚し、誠実に職務を遂行するとともに、職場の秩序の維持に努めなければならない。")

    add_heading(doc, "第7条（遵守事項）", level=2)
    add_para(doc, "従業員は、次の各号に掲げる事項を遵守しなければならない。")
    add_para(doc, "(1) 業務上の指示・命令に従うこと")
    add_para(doc, "(2) 業務上知り得た秘密を漏らさないこと（在職中・退職後を問わない）")
    add_para(doc, "(3) 事業主の許可なく、業務以外の目的で会社の設備・備品を使用しないこと")
    add_para(doc, "(4) 会社の名誉・信用を傷つける行為をしないこと")
    add_para(doc, "(5) ハラスメント（パワハラ・セクハラ等）に該当する行為をしないこと")

    # 第4章 勤務時間・休憩・休日
    add_heading(doc, "第4章　勤務時間・休憩・休日", level=1)

    add_heading(doc, "第8条（始業・終業時刻、休憩時間）", level=2)
    add_para(doc, "1. 1日の所定労働時間は8時間とし、始業・終業時刻および休憩時間は次のとおりとする。")
    add_para(doc, "　始業：午前9時00分　／　終業：午後6時00分　／　休憩：正午〜午後1時（60分）")
    add_para(doc, "2. 業務の都合により、始業・終業時刻を繰り上げ、または繰り下げることがある。")
    add_note(doc, "※実態に合わせて時刻を変更してください。シフト制の場合は別途シフト表で定める旨を記載します。")

    add_heading(doc, "第9条（休日）", level=2)
    add_para(doc, "休日は次のとおりとする。")
    add_para(doc, "(1) 日曜日")
    add_para(doc, "(2) 土曜日")
    add_para(doc, "(3) 国民の祝日（休日法に定めるもの）")
    add_para(doc, "(4) 年末年始（12月29日〜1月3日）")
    add_para(doc, "(5) その他事業主が指定する日")

    add_heading(doc, "第10条（時間外・休日労働）", level=2)
    add_para(doc, "1. 業務の都合により、所定労働時間を超えて、または休日に労働させることがある。")
    add_para(doc, "2. 時間外労働・休日労働を命じる場合は、あらかじめ労使協定（36協定）を締結し、所轄の労働基準監督署長へ届け出る。")
    add_note(doc, "※従業員を雇用する場合、36協定の届出は必須です。提出していない時間外労働は違法となります。")

    # 第5章 休暇
    add_heading(doc, "第5章　休暇", level=1)

    add_heading(doc, "第11条（年次有給休暇）", level=2)
    add_para(doc, "1. 雇入れの日から6か月間継続勤務し、所定労働日の8割以上出勤した従業員には、10日の年次有給休暇を付与する。以後1年経過ごとに、勤続年数に応じて法定の日数を付与する。")
    add_para(doc, "2. 年次有給休暇は、従業員が請求する時季に与える。ただし、事業の正常な運営を妨げる場合は、他の時季に変更することがある。")
    add_para(doc, "3. 年次有給休暇のうち5日については、事業主が時季を指定して取得させることがある（年5日の取得義務）。")

    add_heading(doc, "第12条（産前産後休業・育児休業・介護休業等）", level=2)
    add_para(doc, "従業員は、法令の定めるところにより、産前産後休業、育児休業、介護休業、子の看護休暇、介護休暇を取得することができる。")

    # 第6章 賃金
    add_heading(doc, "第6章　賃金", level=1)

    add_heading(doc, "第13条（賃金の決定・計算・支払い）", level=2)
    add_para(doc, "1. 賃金は、基本給および各種手当により構成し、個別の労働契約書により定める。")
    add_para(doc, "2. 賃金の計算期間は毎月1日から末日までとし、翌月末日に支払う。")
    add_para(doc, "3. 賃金は、所得税・住民税・社会保険料等を控除のうえ、本人が指定する金融機関の口座に振り込む。")
    add_note(doc, "※締め日・支払日は実態に合わせて変更してください。例：「20日締め・当月25日払い」など。")

    add_heading(doc, "第14条（割増賃金）", level=2)
    add_para(doc, "時間外労働・休日労働・深夜労働を行った場合は、労働基準法に定める割増率により計算した割増賃金を支払う。")

    # 第7章 退職・解雇
    add_heading(doc, "第7章　退職・解雇", level=1)

    add_heading(doc, "第15条（退職）", level=2)
    add_para(doc, "従業員が次の各号のいずれかに該当する場合は、退職とする。")
    add_para(doc, "(1) 自己の都合により退職を申し出て、事業主が承認した場合（原則として退職希望日の30日前までに申し出ること）")
    add_para(doc, "(2) 雇用期間の定めがある場合に、その期間が満了した場合")
    add_para(doc, "(3) 死亡した場合")
    add_para(doc, "(4) 定年に達した場合")

    add_heading(doc, "第16条（解雇）", level=2)
    add_para(doc, "従業員が次の各号のいずれかに該当する場合は、解雇することがある。")
    add_para(doc, "(1) 勤務成績または業務能率が著しく不良で、向上の見込みがないと認められたとき")
    add_para(doc, "(2) 心身の故障により業務に堪えられないと認められたとき")
    add_para(doc, "(3) 事業の縮小その他事業の運営上やむを得ない事情があるとき")
    add_para(doc, "(4) その他前各号に準ずるやむを得ない事由があるとき")
    add_note(doc, "※解雇には客観的合理性と社会通念上の相当性が必要です（労働契約法16条）。安易な解雇は無効となるリスクがあります。")

    # 第8章 安全衛生
    add_heading(doc, "第8章　安全衛生", level=1)

    add_heading(doc, "第17条（安全衛生の確保）", level=2)
    add_para(doc, "1. 事業主は、従業員の安全と健康の確保のため、必要な措置を講ずる。")
    add_para(doc, "2. 従業員は、安全衛生に関する法令および事業主の指示を遵守し、自らの健康管理に努めなければならない。")

    add_heading(doc, "第18条（健康診断）", level=2)
    add_para(doc, "事業主は、常時使用する従業員に対して、年1回以上の定期健康診断を実施する。従業員はこれを受診しなければならない。")

    # 第9章 制裁
    add_heading(doc, "第9章　制裁", level=1)

    add_heading(doc, "第19条（制裁の種類）", level=2)
    add_para(doc, "従業員がこの規則に違反したときは、その情状により次の制裁を行う。")
    add_para(doc, "(1) けん責　始末書を提出させ、将来を戒める")
    add_para(doc, "(2) 減給　1回の額が平均賃金の1日分の半額を超えず、総額が一賃金支払期における賃金総額の10分の1を超えない範囲で行う")
    add_para(doc, "(3) 出勤停止　7日以内の出勤停止を命じ、その期間中の賃金は支払わない")
    add_para(doc, "(4) 諭旨退職　退職を勧告し、応じない場合は懲戒解雇とする")
    add_para(doc, "(5) 懲戒解雇　即時に解雇する")

    # 附則
    add_heading(doc, "附則", level=1)
    add_para(doc, "1. この規則は、　　年　　月　　日から施行する。")
    add_para(doc, "2. この規則の改廃は、事業主が従業員代表の意見を聴いたうえで行う。")
    add_para(doc, "3. 改正または廃止した場合は、改正後の規則を従業員に周知する。")

    # 末尾の注意書き
    doc.add_page_break()
    add_heading(doc, "雛形ご利用にあたっての注意事項", level=1)
    add_para(doc, "・本サンプルは、開業したて・個人事業主・副業者・初めて従業員を雇う方を対象とした、最小構成の就業規則雛形です。")
    add_para(doc, "・常時10人以上の労働者を雇用する事業所は、就業規則を作成し、所轄の労働基準監督署長へ届け出る義務があります（労働基準法第89条）。")
    add_para(doc, "・10人未満の事業所でも、就業規則を整備しておくことでトラブル防止・助成金申請などに有利に働きます。")
    add_para(doc, "・実際にご利用される際は、自社の業態・労働実態に合わせて条文を調整してください。")
    add_para(doc, "・条文の追加・削除や、運用に関するご相談は、橋本貴嗣社会保険労務士事務所までお問い合わせください。")

    add_para(doc, "")
    add_para(doc, "──────────────────────")
    add_para(doc, "橋本貴嗣社会保険労務士事務所")
    add_para(doc, "Hashimoto SR & FP1 Office")
    add_para(doc, "https://www.sharoushi-t.com/")
    add_para(doc, "──────────────────────")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Generated: {OUTPUT}")


if __name__ == "__main__":
    main()
